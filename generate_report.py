"""Generate the TaskForge OS project report as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "Times New Roman"


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "000000")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def apply_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def set_paragraph_spacing(p, before=0, after=6, line=1.5):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_heading(doc, text, level=1, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sizes = {1: 16, 2: 14, 3: 13}
    run = p.add_run(text)
    apply_font(run, size=sizes.get(level, 13), bold=True)
    set_paragraph_spacing(p, before=12, after=6, line=1.15)
    return p


def add_para(doc, text, size=12, bold=False, italic=False, justify=True, indent=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    apply_font(run, size=size, bold=bold, italic=italic)
    set_paragraph_spacing(p, before=0, after=6, line=1.5)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.4)
    return p


def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    apply_font(run, size=size)
    set_paragraph_spacing(p, before=0, after=4, line=1.4)
    return p


def add_bullet_kv(doc, key, value, size=12):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"{key}: ")
    apply_font(r1, size=size, bold=True)
    r2 = p.add_run(value)
    apply_font(r2, size=size)
    set_paragraph_spacing(p, before=0, after=4, line=1.4)
    return p


def add_centered_run(doc, text, size=12, bold=False, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    apply_font(run, size=size, bold=bold, italic=italic)
    set_paragraph_spacing(p, before=0, after=space_after, line=1.2)
    return p


def add_screenshot_placeholder(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n[ Screenshot Placeholder ]\n\n")
    apply_font(run, size=12, italic=True, color=RGBColor(0x80, 0x80, 0x80))
    # Pad height
    for _ in range(4):
        pp = cell.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run("")
        apply_font(r, size=12)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cap.add_run(caption)
    apply_font(crun, size=11, italic=True)
    set_paragraph_spacing(cap, before=2, after=10, line=1.15)


def add_code_block(doc, code_text):
    # Render as monospaced-style block via shaded single-cell table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.2)
    set_cell_borders(cell)
    # Light gray shading
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    tc_pr.append(shd)
    cell.paragraphs[0].text = ""
    for line in code_text.splitlines():
        p = cell.add_paragraph()
        run = p.add_run(line if line else " ")
        # Use Consolas for code clarity but keep doc default elsewhere
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), "Consolas")
        set_paragraph_spacing(p, before=0, after=0, line=1.1)
    # remove blank first paragraph
    first = cell.paragraphs[0]
    if first.text.strip() == "":
        first._element.getparent().remove(first._element)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=6, line=1.0)


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def make_table(doc, headers, rows, col_widths=None, header_color="2F5496", header_text_color=RGBColor(0xFF, 0xFF, 0xFF)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_borders(cell)
        shade_cell(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        apply_font(run, size=11, bold=True, color=header_text_color)
        if col_widths:
            cell.width = col_widths[i]

    # Rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].add_run(str(val))
            apply_font(run, size=11)
            if col_widths:
                cell.width = col_widths[c_idx]
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=6, line=1.0)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------- Build the document ----------

doc = Document()

# Set default style font globally
style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = Pt(12)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
    rfonts.set(qn(attr), FONT_NAME)

# Margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ============== TITLE PAGE ==============
add_centered_run(doc, "A PROJECT REPORT", size=16, bold=True, space_after=6)
add_centered_run(doc, "ON", size=13, space_after=18)

add_centered_run(doc, "TaskForge OS", size=22, bold=True, space_after=4)
add_centered_run(doc, "Banking System on a Custom Operating System Kernel", size=15, italic=True, space_after=24)

add_centered_run(doc, "Submitted in partial fulfillment of the requirements", size=12, space_after=2)
add_centered_run(doc, "for the course of", size=12, space_after=2)
add_centered_run(doc, "OPERATING SYSTEMS", size=14, bold=True, space_after=24)

add_centered_run(doc, "Submitted by:", size=12, bold=True, space_after=8)

# Student details table
student_headers = ["Sr. No.", "Roll No.", "Name", "Contact No.", "Email"]
student_rows = [
    ["1", "12414268", "AARUSH ASHISH BAKSHI", "9225515642", "aarush.bakshi24@vit.edu"],
    ["2", "12414364", "ADITYA SAMEER CHIMURKAR", "7385174280", "aditya.chimurkar24@vit.edu"],
    ["3", "12414363", "AYUSH YOGESH AGNIHOTRI", "9075467236", "ayush.agnihotri24@vit.edu"],
    ["4", "12414251", "AMEYA AMIT BORKAR", "8999260116", "ameya.borkar24@vit.edu"],
    ["5", "12413706", "ARNAV ANAND GUPTA", "9112200538", "anand.arnav241@vit.edu"],
]
make_table(
    doc,
    student_headers,
    student_rows,
    col_widths=[Inches(0.6), Inches(1.0), Inches(2.4), Inches(1.2), Inches(2.0)],
)

add_centered_run(doc, "Under the guidance of", size=12, space_after=4)
add_centered_run(doc, "Prof. Amruta Amune", size=14, bold=True, space_after=24)

add_centered_run(doc, "DEPARTMENT OF COMPUTER ENGINEERING", size=12, bold=True, space_after=2)
add_centered_run(doc, "VISHWAKARMA INSTITUTE OF TECHNOLOGY, PUNE", size=12, bold=True, space_after=2)
add_centered_run(doc, "Academic Year 2025–2026", size=12, italic=True, space_after=4)

page_break(doc)

# ============== TABLE OF CONTENTS ==============
add_heading(doc, "TABLE OF CONTENTS", level=1, center=True)

toc_rows = [
    ["1", "Introduction", "3"],
    ["2", "Objectives", "4"],
    ["3", "System Architecture and OS Concept Mapping", "5"],
    ["3.1", "Layered Architecture Overview", "5"],
    ["3.2", "System Call Interface", "6"],
    ["3.3", "OS Concepts Mapping (Unit-wise)", "7"],
    ["3.4", "Worked Example: A Single Banking Operation", "9"],
    ["4", "Implementation Technology", "10"],
    ["5", "Results and Output", "11"],
    ["6", "Conclusion and Future Scope", "13"],
    ["7", "References", "14"],
]
make_table(
    doc,
    ["Sr. No.", "Chapter / Section", "Page No."],
    toc_rows,
    col_widths=[Inches(1.0), Inches(4.5), Inches(1.0)],
)

page_break(doc)

# ============== 1. INTRODUCTION ==============
add_heading(doc, "1. INTRODUCTION", level=1)

add_para(
    doc,
    "An Operating System (OS) is the software layer that manages a computer's hardware and provides "
    "a controlled environment in which applications can run. Concepts such as processes, scheduling, "
    "memory management, deadlocks, file systems, and I/O are fundamental to every modern OS, yet "
    "they are usually presented to students as isolated algorithms — Round Robin in one lecture, the "
    "Banker's Algorithm in another, page replacement in a third. The connection between these "
    "algorithms and a real, working application is rarely demonstrated end-to-end.",
    indent=True,
)

add_para(
    doc,
    "TaskForge OS bridges this gap. It is a complete, self-contained Operating System kernel "
    "implemented in C, exposing a System Call interface on top of which a fully functional Banking "
    "Application is built. Every banking action — opening an account, depositing money, transferring "
    "funds between accounts — is routed through the kernel, where it triggers process creation, "
    "memory allocation, mutual-exclusion locking, deadlock-safety checks, page-cache lookups, file "
    "I/O, and disk scheduling. Each of these stages is logged in real time, so the user can observe "
    "exactly how the OS handles the request.",
    indent=True,
)

add_para(
    doc,
    "The project is delivered in three complementary forms: a native C command-line application, a "
    "native Win32 graphical interface, and a Python/Flask web port that ships with an interactive "
    "dashboard. Together they let the team showcase classical OS theory in action against a "
    "real-world transactional workload.",
    indent=True,
)

page_break(doc)

# ============== 2. OBJECTIVES ==============
add_heading(doc, "2. OBJECTIVES", level=1)

add_para(
    doc,
    "The TaskForge OS project was undertaken with the following objectives, each chosen to "
    "demonstrate a specific set of operating system concepts within a unified application:",
    indent=True,
)

objectives = [
    ("Design a custom OS kernel in C", "Implement a Process Control Block (PCB), scheduler, memory manager, file system, deadlock manager, and I/O subsystem from first principles, without relying on a host OS for any of the simulated abstractions."),
    ("Expose a clean System Call interface", "Provide more than twenty-five system calls (sys_fork, sys_alloc, sys_open, sys_lock, sys_read, sys_write, sys_close, sys_unlock, etc.) so that applications interact with the kernel in the same way they would on a real OS."),
    ("Build a real banking application on top of the kernel", "Support account creation, deposit, withdrawal, transfer, balance enquiry, and statement generation. Persist account data via the kernel's file system layer and serialise transactions through kernel-managed locks."),
    ("Showcase every major OS unit", "Cover process management and concurrency, CPU scheduling (FCFS, SJF, SRTF, Round Robin, Priority), deadlock handling (Banker's Algorithm, Resource Allocation Graph, prevention via resource ordering), memory management (First/Best/Worst/Next Fit, paging, page replacement — FIFO, LRU, Optimal, Clock), and I/O / disk scheduling (FCFS, SSTF, SCAN, C-SCAN)."),
    ("Make the OS observable in real time", "Stream a kernel trace for every banking operation, and present a dashboard with live counters for processes, memory, cache hits/misses, deadlock state, file descriptors, and disk seeks."),
    ("Allow runtime reconfiguration", "Let the user switch the active scheduler, memory allocation strategy, page replacement policy, and disk scheduling algorithm on the fly to compare their behaviour without recompilation."),
    ("Provide multiple front-ends", "Deliver a CLI, a Win32 GUI, and a browser-based SPA so the same kernel can be demonstrated in lab, on a personal machine, and during a viva."),
]
for title, desc in objectives:
    add_bullet_kv(doc, title, desc)

page_break(doc)

# ============== 3. SYSTEM ARCHITECTURE ==============
add_heading(doc, "3. SYSTEM ARCHITECTURE AND OS CONCEPT MAPPING", level=1)

add_heading(doc, "3.1 Layered Architecture Overview", level=2)
add_para(
    doc,
    "TaskForge OS follows the classical three-layer architecture used by real-world operating "
    "systems. The Banking Application sits at the top and never accesses any kernel data structure "
    "directly. All requests cross a well-defined System Call interface, which is the only contract "
    "between user code and the kernel. Below that interface, the kernel implements the six "
    "subsystems traditionally associated with an OS: process management, scheduling, memory "
    "management, file system, deadlock management, and I/O. The diagram below captures this layout.",
    indent=True,
)

add_code_block(
    doc,
    "+--------------------------------------------------+\n"
    "|  BANKING APPLICATION              (app/bank.c)   |\n"
    "|  Create Account | Deposit | Withdraw             |\n"
    "|  Transfer       | Balance | Statement            |\n"
    "+--------------------------------------------------+\n"
    "|  SYSTEM CALL API               (os_syscall.h)    |\n"
    "|  sys_fork  sys_alloc  sys_open  sys_lock         |\n"
    "|  sys_read  sys_write  sys_close sys_unlock       |\n"
    "+--------------------------------------------------+\n"
    "|  TASKFORGE OS KERNEL          (os_kernel.c)      |\n"
    "|  Process Mgr | Scheduler | Memory Mgr            |\n"
    "|  File System | Deadlock Mgr | I/O Subsystem      |\n"
    "+--------------------------------------------------+",
)

add_heading(doc, "3.2 System Call Interface", level=2)
add_para(
    doc,
    "The system call layer (os_syscall.h) defines the boundary between unprivileged application "
    "code and privileged kernel routines. Each call accepts a process identifier and the resources "
    "the caller needs, returning a status code. The most commonly used calls in the Banking "
    "Application are:",
    indent=True,
)
add_bullet(doc, "sys_fork() — create a new kernel process to handle a banking transaction.")
add_bullet(doc, "sys_alloc() / sys_free() — allocate and release memory using the configured strategy.")
add_bullet(doc, "sys_lock() / sys_unlock() — acquire and release per-account mutual exclusion locks, with Banker's-algorithm safety checks.")
add_bullet(doc, "sys_open() / sys_read() / sys_write() / sys_close() — file operations that are ultimately serviced by the disk-scheduling subsystem.")
add_bullet(doc, "sys_yield() / sys_exit() — voluntary CPU release and process termination.")

add_heading(doc, "3.3 OS Concepts Mapping (Unit-wise)", level=2)
add_para(
    doc,
    "The table below maps each unit of the Operating Systems syllabus to the file, function, and "
    "kernel feature that implements it. This mapping is also exposed live in the web dashboard, "
    "where each row is paired with a metric that updates as the user performs banking operations.",
    indent=True,
)

mapping_rows = [
    ["I", "System Calls, OS Types",
     "os_syscall.h — 25+ system calls (sys_fork, sys_alloc, sys_open, sys_lock, …)"],
    ["II", "Process Management, Threads, Concurrency",
     "PCB with five-state model; pthreads, mutexes, semaphores; Producer–Consumer, Readers–Writers, Dining Philosophers demos"],
    ["III", "CPU Scheduling",
     "FCFS, SJF, SRTF, Round Robin, Priority (preemptive and non-preemptive); Gantt charts and comparison mode"],
    ["IV", "Deadlocks",
     "Banker's Algorithm for avoidance, Resource Allocation Graph, deadlock detection, resource ordering for prevention"],
    ["V", "Memory Management",
     "First / Best / Worst / Next Fit allocation; paging, segmentation, TLB simulation; FIFO, LRU, Optimal, Clock page replacement"],
    ["VI", "I/O and File Management",
     "Disk scheduling (FCFS, SSTF, SCAN, C-SCAN), virtual file system, I/O buffering, free-space management"],
]
make_table(
    doc,
    ["Unit", "Topic", "Implementation in TaskForge OS"],
    mapping_rows,
    col_widths=[Inches(0.6), Inches(2.0), Inches(3.9)],
)

add_heading(doc, "3.4 Worked Example: A Single Banking Operation", level=2)
add_para(
    doc,
    "To make the concept mapping concrete, consider what happens when a user deposits ₹500 into "
    "Account #1. The kernel trace below is reproduced verbatim from a live run of TaskForge OS, "
    "and shows seven distinct OS subsystems collaborating on a single user-level action.",
    indent=True,
)

add_code_block(
    doc,
    "[BANK] Depositing $500.00 to Account #1...\n"
    "[OS]   Process P5 'deposit' created (Priority: 4)\n"
    "[OS]   P5 state: NEW -> READY -> RUNNING\n"
    "[OS]   Memory allocated: 1 KB at address 128 (Best Fit)\n"
    "[OS]   Resource lock on Account #1: GRANTED (Banker's: SAFE)\n"
    "[OS]   Cache access page 1: HIT\n"
    "[OS]   File write: 24 bytes to acc_001.dat\n"
    "[OS]   Disk I/O: track 13, seek distance 37 (SCAN)\n"
    "[OS]   Resource unlocked: Account #1\n"
    "[OS]   Memory freed for P5\n"
    "[OS]   P5 state: RUNNING -> TERMINATED\n"
    "[BANK] Deposit complete. New balance: $1,500.00",
)

add_para(
    doc,
    "A transfer operation goes one step further: it requires locks on two accounts simultaneously, "
    "and would deadlock under naive locking. TaskForge OS prevents this by enforcing resource "
    "ordering — the lower-numbered account is always locked first — and additionally validates the "
    "request through the Banker's Algorithm before any resources are committed. Together these "
    "techniques demonstrate both deadlock prevention and deadlock avoidance on the same code path.",
    indent=True,
)

page_break(doc)

# ============== 4. IMPLEMENTATION TECHNOLOGY ==============
add_heading(doc, "4. IMPLEMENTATION TECHNOLOGY", level=1)

add_para(
    doc,
    "TaskForge OS has been engineered to run on commodity student hardware while still exhibiting "
    "the breadth of a full operating system. The choice of technology was driven by three "
    "constraints: portability across Windows and Linux lab machines, the ability to use real "
    "operating-system primitives (rather than fake ones), and a low barrier to entry for the "
    "evaluator running the demo.",
    indent=True,
)

tech_rows = [
    ["Programming Language (Kernel & App)", "C (C11 standard) — for the OS kernel, system call layer, banking application, and Win32 GUI"],
    ["Programming Language (Web Port)", "Python 3.8+ with Flask — for the web dashboard and REST API that exposes the kernel state"],
    ["Threading Library", "POSIX threads (pthreads) — used for kernel processes and concurrency demos (Producer–Consumer, Readers–Writers, Dining Philosophers)"],
    ["Synchronisation Primitives", "pthread mutexes and semaphores for locking; spin-waits for short critical sections; condition variables for blocking"],
    ["GUI (Native)", "Win32 API with Common Controls — used for the desktop GUI variant of the project"],
    ["GUI (Web)", "HTML5 + CSS3 + vanilla JavaScript single-page application served by Flask, with a dark theme and live polling of kernel metrics"],
    ["Build System", "GCC via MinGW / MSYS2 on Windows (and stock GCC on Linux); a build.bat script and a GNU Makefile are both provided"],
    ["Persistence", "Plain-file persistence via the kernel's virtual file system — each account is stored as acc_XXX.dat and accessed through sys_open/read/write"],
    ["External Dependencies", "Standard C library, pthreads, Win32 API, and Flask. No third-party frameworks, databases, or runtime services are required."],
    ["Target Platforms", "Windows 10 / 11 and any Linux distribution with GCC and Python 3 installed"],
]
make_table(
    doc,
    ["Component", "Technology / Library"],
    tech_rows,
    col_widths=[Inches(2.0), Inches(4.5)],
)

add_para(
    doc,
    "The codebase is organised into clearly separated layers. The kernel/ directory contains the "
    "kernel headers and implementation; app/ holds the banking application; main_v2.c is the boot "
    "and menu entry point; gui/ contains the Win32 desktop interface; and web/ contains the Flask "
    "port together with its single-page-application front-end. A bonus v1 simulator under src/ "
    "and include/ provides standalone interactive demos of every algorithm in the syllabus.",
    indent=True,
)

page_break(doc)

# ============== 5. RESULTS AND OUTPUT ==============
add_heading(doc, "5. RESULTS AND OUTPUT", level=1)

add_para(
    doc,
    "The screenshots below illustrate TaskForge OS in operation across its three front-ends. Each "
    "placeholder is to be replaced with the corresponding capture from a live demo run.",
    indent=True,
)

add_heading(doc, "5.1 Web Dashboard — Banking Tab", level=2)
add_screenshot_placeholder(doc, "Figure 5.1: Banking tab showing account creation, deposit, withdraw, and transfer with the live kernel trace.")

add_heading(doc, "5.2 Web Dashboard — OS Dashboard Tab", level=2)
add_screenshot_placeholder(doc, "Figure 5.2: Live state of processes, memory, page cache, deadlock manager, file system, and I/O subsystem.")

add_heading(doc, "5.3 Web Dashboard — OS Concepts Tab", level=2)
add_screenshot_placeholder(doc, "Figure 5.3: Each syllabus unit mapped to the implementation file/function and a live kernel metric.")

add_heading(doc, "5.4 Web Dashboard — Configuration Tab", level=2)
add_screenshot_placeholder(doc, "Figure 5.4: Switching scheduler, memory strategy, cache policy, and disk algorithm at runtime.")

add_heading(doc, "5.5 Native Win32 GUI", level=2)
add_screenshot_placeholder(doc, "Figure 5.5: Native desktop interface (taskforge_gui_v2.exe) running the banking application with an embedded OS dashboard.")

add_heading(doc, "5.6 Native CLI", level=2)
add_screenshot_placeholder(doc, "Figure 5.6: Command-line build (taskforge_v2.exe) producing the kernel trace for a deposit operation.")

add_heading(doc, "5.7 v1 Algorithm Simulator", level=2)
add_screenshot_placeholder(doc, "Figure 5.7: Standalone simulator demonstrating CPU scheduling, page replacement, and disk scheduling with user-supplied inputs.")

page_break(doc)

# ============== 6. CONCLUSION AND FUTURE SCOPE ==============
add_heading(doc, "6. CONCLUSION AND FUTURE SCOPE", level=1)

add_heading(doc, "6.1 Conclusion", level=2)
add_para(
    doc,
    "TaskForge OS demonstrates that the entire Operating Systems syllabus can be unified into a "
    "single working system rather than a collection of disjoint algorithm demos. By implementing a "
    "real OS kernel in C, exposing it through a system call interface, and building a banking "
    "application that exercises every kernel subsystem on every transaction, the project shows in a "
    "very tangible way how processes, scheduling, memory, deadlocks, file systems, and I/O "
    "co-operate to service even the simplest user request. The accompanying live dashboard and "
    "configurable algorithms turn the project into a teaching aid: students can observe the effect "
    "of switching from FIFO to LRU page replacement, or from FCFS to SCAN disk scheduling, on a "
    "real workload rather than a contrived input. The three front-ends — CLI, native Win32 GUI, and "
    "browser-based SPA — make the system suitable for lab use, viva-voce demonstration, and "
    "self-study alike.",
    indent=True,
)

add_heading(doc, "6.2 Future Scope", level=2)
add_para(
    doc,
    "While the current version satisfies all the goals set at the start of the project, several "
    "enhancements would meaningfully extend its educational and practical value:",
    indent=True,
)
future = [
    ("Networking subsystem", "Add a simple network stack with sockets so that banking transactions can be performed remotely, exposing concepts such as sockets, message queues, and inter-process communication."),
    ("Multi-core scheduling", "Extend the scheduler to manage multiple logical CPUs with load balancing, demonstrating SMP scheduling and processor affinity."),
    ("Persistent transaction log", "Introduce write-ahead logging and crash recovery so that banking state survives unexpected termination, illustrating journaling file-system techniques."),
    ("Security and authentication", "Add user accounts, password hashing, and per-process privilege levels to demonstrate access control and protection mechanisms."),
    ("Virtual memory with demand paging", "Move from a flat physical-address allocator to a paged virtual-address space backed by a swap file, with page faults serviced by the I/O subsystem."),
    ("Containerisation of the demo", "Ship a Docker image that boots the web dashboard automatically, lowering the barrier to evaluation by external reviewers."),
    ("Mobile-friendly dashboard", "Make the Flask front-end responsive so that the OS dashboard can be observed from a phone or tablet during a presentation."),
    ("Automated test suite", "Add unit tests for each scheduling, memory, and disk algorithm, plus integration tests that drive the banking application through scripted workloads."),
]
for title, desc in future:
    add_bullet_kv(doc, title, desc)

page_break(doc)

# ============== 7. REFERENCES ==============
add_heading(doc, "7. REFERENCES", level=1)

refs = [
    "Silberschatz, A., Galvin, P. B., and Gagne, G., Operating System Concepts, 10th Edition, Wiley, 2018.",
    "Tanenbaum, A. S., and Bos, H., Modern Operating Systems, 4th Edition, Pearson, 2014.",
    "Stallings, W., Operating Systems: Internals and Design Principles, 9th Edition, Pearson, 2018.",
    "Dhamdhere, D. M., Operating Systems: A Concept-Based Approach, 3rd Edition, McGraw-Hill, 2012.",
    "Kerrisk, M., The Linux Programming Interface, No Starch Press, 2010.",
    "POSIX Threads Programming — Lawrence Livermore National Laboratory, https://hpc-tutorials.llnl.gov/posix/",
    "GNU C Library (glibc) Reference Manual, https://www.gnu.org/software/libc/manual/",
    "Microsoft Windows API Reference — Win32 Programming, https://learn.microsoft.com/en-us/windows/win32/api/",
    "Flask Documentation, Pallets Projects, https://flask.palletsprojects.com/",
    "Python 3 Language Reference, Python Software Foundation, https://docs.python.org/3/",
    "GCC, the GNU Compiler Collection, https://gcc.gnu.org/onlinedocs/",
    "Operating Systems syllabus and reference material provided by the Department of Computer Engineering, Vishwakarma Institute of Technology, Pune.",
]
for i, ref in enumerate(refs, start=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"[{i}]  ")
    apply_font(r, size=12, bold=True)
    r2 = p.add_run(ref)
    apply_font(r2, size=12)
    set_paragraph_spacing(p, before=0, after=6, line=1.4)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)

# Save
output_path = r"C:\Users\ameya\Documents\TaskForge\TaskForge_OS_Project_Report.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
